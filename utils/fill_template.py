from docx import Document

def fill_placeholders(template_path, replacements, output_path):
    doc = Document(template_path)

    def replace_text(text):
        for key, value in replacements.items():
            text = text.replace(f"{{{key}}}", value)
        return text

    for para in doc.paragraphs:
        para.text = replace_text(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = replace_text(cell.text)

    doc.save(output_path)
